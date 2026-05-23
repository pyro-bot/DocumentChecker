import sys
import os
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from pylatex import Document as LatexDocument
from pylatex import Command, NoEscape

from .text import parse_paragraphs, parse_list_paragraphs
from .tables import parse_table
from .images import extract_images, get_image_rel_ids



def iter_block_items(parent):
    from docx.document import Document as DocxDocType
    if isinstance(parent, DocxDocType):
        parent_elm = parent.element.body

    else:
        raise ValueError("parent must be a Document or _Cell")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def convert_to_latex(docx_path, output_path, image_dir='documents/output/images'):
    document = DocxDocument(docx_path)

    image = extract_images(document, image_dir)

    latex_doc = LatexDocument()

    latex_doc.preamble.append(Command('usepackage', 'babel', options=['russian']))
    latex_doc.preamble.append(Command('usepackage', 'fontenc'))
    latex_doc.preamble.append(Command('usepackage', 'graphicx'))
    latex_doc.preamble.append(Command('usepackage', 'float'))
    latex_doc.preamble.append(Command('usepackage', 'longtable'))
    latex_doc.preamble.append(Command('usepackage', 'array'))


    flag_itemize = False
    flag_enumerate = False

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            paragraphs = block
            img_rel_ids = get_image_rel_ids(paragraphs)
            if img_rel_ids:
                if flag_itemize:
                    latex_doc.append(NoEscape(r'\end{itemize}'))
                    flag_itemize = False
                if flag_enumerate:
                    latex_doc.append(NoEscape(r'\end{enumerate}'))
                    flag_enumerate = False

                for rel_id in img_rel_ids:
                    if rel_id in image:
                        img_path = os.path.join(image_dir, os.path.basename(image[rel_id]))
                        safe_path = img_path.replace('\\', '/').replace('_', r'\_')
                        latex_doc.append(NoEscape(r'\begin{figure}[H]'))
                        latex_doc.append(NoEscape(r'\centering'))
                        latex_doc.append(NoEscape(rf'\includegraphics[width=0.8\textwidth]{{{safe_path}}}'))
                        latex_doc.append(NoEscape(r'\end{figure}'))
                continue

            if paragraphs.style.name in ['List Paragraph', 'List Number', 'List Bullet']:
                flag_itemize, flag_enumerate = parse_list_paragraphs(paragraphs, latex_doc, flag_itemize, flag_enumerate)
            else:
                flag_itemize, flag_enumerate = parse_paragraphs(paragraphs, latex_doc, flag_itemize, flag_enumerate)
        elif isinstance(block, Table):
            if flag_itemize:
                latex_doc.append(NoEscape(r'\end{itemize}'))
                flag_itemize = False
            if flag_enumerate:
                latex_doc.append(NoEscape(r'\end{enumerate}'))
                flag_enumerate = False

            parse_table(block, latex_doc)


    if flag_itemize:
        latex_doc.append(NoEscape(r'\end{itemize}'))

    if flag_enumerate:
        latex_doc.append(NoEscape(r'\end{enumerate}'))

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(latex_doc.dumps())


def main():
    input_path = sys.argv[1]
    output_path = sys.argv[2]

    convert_to_latex(input_path, output_path)
    print(f"Результат сохранен в {output_path}")


if __name__ == "__main__":
    main()
